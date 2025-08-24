#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
创建一个包含加点字的测试docx文件，用于验证预处理器
"""

from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn

def create_test_docx_with_dots():
    """创建包含加点字的测试docx文件"""
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('语文试题测试 - 加点字识别', 0)
    
    # 添加材料段落（包含拼音标注）
    material_para = doc.add_paragraph()
    material_para.add_run('资料：经')
    
    # 添加带拼音的"筹"字
    run_chou = material_para.add_run('筹')
    run_chou.add_text('（chóu）划，1949年2月，全国新华书店第一届出版工作会议在北京')
    
    # 添加带拼音的"召"字  
    run_zhao = material_para.add_run('召')
    run_zhao.add_text('（zhāo）开。北京新华书店始终传承红色基因，')
    
    # 添加带拼音的"砥"字
    run_di = material_para.add_run('砥')
    run_di.add_text('（dǐ）砺"新华精神"，为广大读者')
    
    # 添加带拼音的"提"字
    run_ti = material_para.add_run('提')  
    run_ti.add_text('（tí）供科学文化知识。')
    
    # 手动添加加点字格式到特定字符
    try:
        # 为"筹"字添加加点格式
        chou_xml = run_chou._element
        rPr = chou_xml.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            chou_xml.insert(0, rPr)
        
        em_element = OxmlElement('w:em')
        em_element.set(qn('w:val'), 'dot')
        rPr.append(em_element)
        print("✅ 为'筹'字添加了加点格式")
        
        # 为其他字符添加类似格式
        for char, run in [('召', run_zhao), ('砥', run_di), ('提', run_ti)]:
            char_xml = run._element
            rPr = char_xml.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                char_xml.insert(0, rPr)
            
            em_element = OxmlElement('w:em')
            em_element.set(qn('w:val'), 'dot')
            rPr.append(em_element)
            print(f"✅ 为'{char}'字添加了加点格式")
            
    except Exception as e:
        print(f"⚠️ 添加加点格式时出错: {e}")
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加题目
    question_para = doc.add_paragraph()
    question_para.add_run('2. 你审核资料中标注的字音。下列加点字读音标注不正确的一项是（   ）')
    
    # 添加选项
    options_para = doc.add_paragraph()
    options_para.add_run('A. 筹划    B. 召开    C. 砥砺    D. 提供')
    
    # 保存文档
    output_path = "test_dot_chars.docx"
    doc.save(output_path)
    
    print(f"📁 测试文档已创建: {output_path}")
    return output_path

if __name__ == "__main__":
    print("=" * 50)
    print("创建加点字测试文档")  
    print("=" * 50)
    
    test_file = create_test_docx_with_dots()
    
    print(f"\n💡 接下来可以:")
    print(f"1. 运行: python3 docx_dot_preprocessor.py {test_file}")
    print(f"2. 检查预处理效果")
    print("=" * 50)
