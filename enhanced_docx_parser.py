#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
增强版DOCX文档解析器
专门用于识别和处理Word文档中的特殊格式，包括：
1. 各种下划线样式（波浪线、点状线等）
2. 上标、下标文本
3. 删除线文本
4. 加点字（着重号）
5. 文本颜色和背景色

使用方法：
1. 安装依赖：pip install python-docx lxml
2. 运行脚本处理Word文档
"""

import os
import json
import zipfile
import tempfile
from datetime import datetime
from pathlib import Path
from collections import defaultdict

try:
    from docx import Document
    from docx.enum.text import WD_UNDERLINE
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
except ImportError:
    print("❌ 请安装python-docx库: pip install python-docx")
    exit(1)

try:
    from lxml import etree
except ImportError:
    print("❌ 请安装lxml库: pip install lxml")
    exit(1)

class EnhancedDocxParser:
    def __init__(self):
        # 下划线样式映射 - 只使用确实存在的枚举值
        self.underline_styles = {}
        
        # 安全地添加下划线样式
        styles_to_add = [
            (getattr(WD_UNDERLINE, 'SINGLE', None), "单下划线"),
            (getattr(WD_UNDERLINE, 'DOUBLE', None), "双下划线"),
            (getattr(WD_UNDERLINE, 'THICK', None), "粗下划线"),
            (getattr(WD_UNDERLINE, 'DOTTED', None), "点状下划线"),
            (getattr(WD_UNDERLINE, 'DASH', None), "虚线下划线"),
            (getattr(WD_UNDERLINE, 'DOT_DASH', None), "点划线下划线"),
            (getattr(WD_UNDERLINE, 'DOT_DOT_DASH', None), "点点划线下划线"),
            (getattr(WD_UNDERLINE, 'WAVY', None), "波浪线下划线"),
            (getattr(WD_UNDERLINE, 'DOTTED_HEAVY', None), "粗点状下划线"),
            (getattr(WD_UNDERLINE, 'DASH_HEAVY', None), "粗虚线下划线"),
            (getattr(WD_UNDERLINE, 'WAVY_HEAVY', None), "粗波浪线下划线"),
            (getattr(WD_UNDERLINE, 'WAVY_DOUBLE', None), "双波浪线下划线")
        ]
        
        for style_enum, style_name in styles_to_add:
            if style_enum is not None:
                self.underline_styles[style_enum] = style_name
        
        # 用于存储特殊格式的文本
        self.special_formatted_text = []
        
    def extract_text_with_formatting(self, docx_path):
        """提取带格式的文本内容"""
        print(f"🔍 开始解析文档格式: {docx_path}")
        
        if not os.path.exists(docx_path):
            print(f"❌ 文件不存在: {docx_path}")
            return None
            
        try:
            doc = Document(docx_path)
            all_content = []
            paragraph_count = 0
            
            for para in doc.paragraphs:
                paragraph_count += 1
                para_content = self._analyze_paragraph(para, paragraph_count)
                if para_content['text'].strip():  # 只添加非空段落
                    all_content.append(para_content)
            
            # 提取表格中的内容
            table_count = 0
            for table in doc.tables:
                table_count += 1
                table_content = self._analyze_table(table, table_count)
                all_content.append(table_content)
            
            print(f"✅ 解析完成: {paragraph_count}个段落, {table_count}个表格")
            return {
                'document_path': docx_path,
                'paragraphs': all_content,
                'special_formats_summary': self._summarize_special_formats(),
                'extraction_time': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
        except Exception as e:
            print(f"❌ 解析失败: {e}")
            return None
    
    def _analyze_paragraph(self, paragraph, para_index):
        """分析段落中的格式"""
        para_info = {
            'paragraph_index': para_index,
            'text': paragraph.text,
            'runs': [],
            'special_formats': []
        }
        
        for run_index, run in enumerate(paragraph.runs):
            run_info = self._analyze_run(run, para_index, run_index)
            para_info['runs'].append(run_info)
            
            # 收集特殊格式
            if run_info['formats']:
                para_info['special_formats'].append({
                    'run_index': run_index,
                    'text': run.text,
                    'formats': run_info['formats']
                })
        
        return para_info
    
    def _analyze_table(self, table, table_index):
        """分析表格中的格式"""
        table_info = {
            'type': 'table',
            'table_index': table_index,
            'rows': []
        }
        
        for row_index, row in enumerate(table.rows):
            row_info = {
                'row_index': row_index,
                'cells': []
            }
            
            for cell_index, cell in enumerate(row.cells):
                cell_content = []
                for para in cell.paragraphs:
                    para_content = self._analyze_paragraph(para, f"表格{table_index}-行{row_index}-列{cell_index}")
                    cell_content.append(para_content)
                
                row_info['cells'].append({
                    'cell_index': cell_index,
                    'content': cell_content
                })
            
            table_info['rows'].append(row_info)
        
        return table_info
    
    def _analyze_run(self, run, para_index, run_index):
        """分析文本片段的格式"""
        formats = []
        
        # 检查字体格式
        font = run.font
        
        # 下划线检查
        if font.underline:
            underline_style = self.underline_styles.get(font.underline, f"未知下划线样式({font.underline})")
            formats.append(f"下划线: {underline_style}")
            
            # 特别标记波浪线和点状线 - 安全检查
            wavy_styles = [style for style in [
                getattr(WD_UNDERLINE, 'WAVY', None),
                getattr(WD_UNDERLINE, 'WAVY_HEAVY', None),
                getattr(WD_UNDERLINE, 'WAVY_DOUBLE', None)
            ] if style is not None]
            
            dotted_styles = [style for style in [
                getattr(WD_UNDERLINE, 'DOTTED', None),
                getattr(WD_UNDERLINE, 'DOTTED_HEAVY', None)
            ] if style is not None]
            
            if font.underline in wavy_styles:
                formats.append("⚠️ 波浪线格式")
            elif font.underline in dotted_styles:
                formats.append("⚠️ 点状线格式")
        
        # 上标下标
        if font.superscript:
            formats.append("上标")
        if font.subscript:
            formats.append("下标")
        
        # 删除线
        if font.strike:
            formats.append("删除线")
        
        # 粗体斜体
        if font.bold:
            formats.append("粗体")
        if font.italic:
            formats.append("斜体")
        
        # 字体颜色
        if font.color and font.color.rgb:
            try:
                # 安全地处理颜色值
                rgb_val = font.color.rgb
                if hasattr(rgb_val, '__int__'):
                    color_hex = f"#{int(rgb_val):06x}"
                else:
                    color_hex = str(rgb_val)
                formats.append(f"字体颜色: {color_hex}")
            except Exception as e:
                formats.append(f"字体颜色: 无法解析({e})")
        
        # 字体大小
        if font.size:
            size_pt = font.size.pt
            formats.append(f"字号: {size_pt}磅")
        
        # 字体名称
        if font.name:
            formats.append(f"字体: {font.name}")
        
        # 检查着重号（通过XML分析）
        emphasis_mark = self._check_emphasis_mark(run)
        if emphasis_mark:
            formats.append(f"着重号: {emphasis_mark}")
        
        # 保存特殊格式的文本
        if formats and run.text.strip():
            self.special_formatted_text.append({
                'text': run.text,
                'paragraph': para_index,
                'run': run_index,
                'formats': formats
            })
        
        return {
            'text': run.text,
            'formats': formats
        }
    
    def _check_emphasis_mark(self, run):
        """检查着重号（加点字）"""
        try:
            # 通过XML检查w:em元素
            run_xml = run._element
            em_elements = run_xml.xpath('.//w:em', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            if em_elements:
                em_val = em_elements[0].get(qn('w:val'))
                emphasis_types = {
                    'dot': '点',
                    'comma': '逗号',
                    'circle': '圆圈',
                    'underDot': '下点'
                }
                return emphasis_types.get(em_val, f'着重号({em_val})')
        except:
            pass
        
        return None
    
    def _summarize_special_formats(self):
        """总结特殊格式统计"""
        format_stats = defaultdict(int)
        
        for item in self.special_formatted_text:
            for fmt in item['formats']:
                format_stats[fmt] += 1
        
        return dict(format_stats)
    
    def save_analysis_result(self, analysis_result, output_dir="enhanced_analysis"):
        """保存分析结果"""
        if not analysis_result:
            return None
        
        # 创建输出目录
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        # 生成文件名
        doc_name = Path(analysis_result['document_path']).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # 保存详细分析结果
        detail_file = output_path / f"detailed_analysis_{doc_name}_{timestamp}.json"
        with open(detail_file, 'w', encoding='utf-8') as f:
            json.dump(analysis_result, f, ensure_ascii=False, indent=2)
        print(f"📄 详细分析结果已保存: {detail_file}")
        
        # 保存特殊格式摘要
        summary_file = output_path / f"special_formats_summary_{doc_name}_{timestamp}.txt"
        with open(summary_file, 'w', encoding='utf-8') as f:
            f.write(f"特殊格式分析报告\n")
            f.write(f"文档: {analysis_result['document_path']}\n")
            f.write(f"分析时间: {analysis_result['extraction_time']}\n")
            f.write(f"=" * 50 + "\n\n")
            
            f.write("特殊格式统计:\n")
            for fmt, count in analysis_result['special_formats_summary'].items():
                f.write(f"  {fmt}: {count}次\n")
            
            f.write(f"\n特殊格式文本详情:\n")
            for item in self.special_formatted_text:
                f.write(f"\n段落{item['paragraph']}-片段{item['run']}: \"{item['text']}\"\n")
                for fmt in item['formats']:
                    f.write(f"  └─ {fmt}\n")
        
        print(f"📋 特殊格式摘要已保存: {summary_file}")
        
        return {
            'detail_file': detail_file,
            'summary_file': summary_file
        }
    
    def extract_special_format_text_only(self, docx_path):
        """仅提取包含特殊格式的文本"""
        analysis = self.extract_text_with_formatting(docx_path)
        if not analysis:
            return None
        
        special_texts = []
        for item in self.special_formatted_text:
            special_texts.append({
                'text': item['text'],
                'location': f"段落{item['paragraph']}-片段{item['run']}",
                'formats': item['formats']
            })
        
        return special_texts

def main():
    """主函数示例"""
    import sys
    
    # 默认处理的文档
    if len(sys.argv) > 1:
        docx_file = sys.argv[1]
    else:
        docx_file = "Chinese/精品解析：2025年甘肃省兰州市中考语文真题（解析版）.docx"
    
    # 创建解析器
    parser = EnhancedDocxParser()
    
    # 分析文档
    print("🚀 启动增强版DOCX解析器...")
    analysis_result = parser.extract_text_with_formatting(docx_file)
    
    if analysis_result:
        # 保存结果
        saved_files = parser.save_analysis_result(analysis_result)
        
        # 打印摘要
        print("\n📊 特殊格式统计:")
        for fmt, count in analysis_result['special_formats_summary'].items():
            print(f"  {fmt}: {count}次")
        
        # 显示包含特殊格式的文本示例
        print(f"\n🎯 包含特殊格式的文本示例（前5个）:")
        special_texts = parser.extract_special_format_text_only(docx_file)
        for i, item in enumerate(special_texts[:5]):
            print(f"\n{i+1}. 位置: {item['location']}")
            print(f"   文本: \"{item['text'][:50]}{'...' if len(item['text']) > 50 else ''}\"")
            print(f"   格式: {', '.join(item['formats'])}")
        
        if len(special_texts) > 5:
            print(f"\n   ... 还有 {len(special_texts) - 5} 个包含特殊格式的文本")
        
        print(f"\n✅ 解析完成! 共发现 {len(special_texts)} 个包含特殊格式的文本片段")
    else:
        print("❌ 解析失败")

if __name__ == "__main__":
    main()
