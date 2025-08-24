#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Pandoc Word文档处理工具 - 增强版 (支持特殊格式识别)

使用pandoc将Word文档转换为模型可读的纯文本内容，支持：
1. 文档文本转换 (Pandoc)
2. 图片提取和内容分析 (LLM Vision)
3. 水印文字替换
4. 大模型API调用 (文档结构解析)
5. 🆕 特殊格式识别 (波浪线、下划线、上标下标等)
6. 🆕 着重号检测 (加点字)
7. 🆕 字体样式分析 (颜色、字体、字号)
8. 🆕 格式统计报告

依赖安装：
1. 确保系统已安装pandoc: https://pandoc.org/installing.html
2. 安装python-docx: pip install python-docx
3. 安装其他依赖: pip install pillow requests

使用方法：
1. 运行脚本处理Word文档
2. 自动提取图片并使用LLM分析内容
3. 识别文档中的特殊格式（波浪线、下划线等）
4. 生成包含格式信息的解析结果
5. 保存格式分析报告
"""

import subprocess
import requests
import json
import time
import os
import re
import zipfile
import tempfile
from datetime import datetime
from pathlib import Path
from PIL import Image
import base64
from io import BytesIO
from collections import defaultdict

# 新增：特殊格式识别依赖
try:
    from docx import Document
    from docx.enum.text import WD_UNDERLINE
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
    print("✅ python-docx库可用，支持特殊格式识别")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx库不可用，将跳过特殊格式识别功能")
    print("   安装命令: pip install python-docx")

class PandocWordProcessor:
    def __init__(self):
        self.api_key = "baf9ea42-7e17-4df6-9a22-90127ac8220e"
        self.base_url = "https://ark.cn-beijing.volces.com/api"
        
        # 检查pandoc是否可用
        self.pandoc_available = self._check_pandoc()
        if not self.pandoc_available:
            print("⚠️ 警告: pandoc未安装或不在PATH中")
            print("请访问 https://pandoc.org/installing.html 安装pandoc")
        
        # 新增：特殊格式识别功能初始化
        self.format_detection_enabled = DOCX_AVAILABLE
        self.special_formatted_text = []
        self.format_statistics = defaultdict(int)
        
        if self.format_detection_enabled:
            self._init_format_styles()
    
    def _check_pandoc(self):
        """检查pandoc是否可用"""
        try:
            result = subprocess.run(['pandoc', '--version'], 
                                  capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                print(f"✅ Pandoc可用: {result.stdout.split()[1]}")
                return True
            else:
                print(f"❌ Pandoc检查失败: {result.stderr}")
                return False
        except FileNotFoundError:
            print("❌ 未找到pandoc命令")
            return False
        except subprocess.TimeoutExpired:
            print("❌ Pandoc检查超时")
            return False
        except Exception as e:
            print(f"❌ Pandoc检查异常: {e}")
            return False
    
    def _init_format_styles(self):
        """初始化格式样式映射"""
        if not DOCX_AVAILABLE:
            return
        
        # 下划线样式映射 - 安全地添加下划线样式
        self.underline_styles = {}
        styles_to_add = [
            (getattr(WD_UNDERLINE, 'SINGLE', None), "单下划线"),
            (getattr(WD_UNDERLINE, 'DOUBLE', None), "双下划线"),
            (getattr(WD_UNDERLINE, 'THICK', None), "粗下划线"),
            (getattr(WD_UNDERLINE, 'DOTTED', None), "点状下划线"),
            (getattr(WD_UNDERLINE, 'DASH', None), "虚线下划线"),
            (getattr(WD_UNDERLINE, 'DOT_DASH', None), "点划线下划线"),
            (getattr(WD_UNDERLINE, 'DOT_DOT_DASH', None), "点点划线下划线"),
            (getattr(WD_UNDERLINE, 'WAVY', None), "波浪线下划线"),
            (getattr(WD_UNDERLINE, 'DOTTED_HEAVY', None), "粗点状下划线"),
            (getattr(WD_UNDERLINE, 'DASH_HEAVY', None), "粗虚线下划线"),
            (getattr(WD_UNDERLINE, 'WAVY_HEAVY', None), "粗波浪线下划线"),
            (getattr(WD_UNDERLINE, 'WAVY_DOUBLE', None), "双波浪线下划线")
        ]
        
        for style_enum, style_name in styles_to_add:
            if style_enum is not None:
                self.underline_styles[style_enum] = style_name
        
        print(f"📋 初始化了 {len(self.underline_styles)} 种下划线样式识别")
    
    def _analyze_text_formatting(self, run, para_index=0, run_index=0):
        """分析文本片段的格式"""
        if not DOCX_AVAILABLE:
            return []
        
        formats = []
        font = run.font
        
        # 下划线检查
        if font.underline:
            underline_style = self.underline_styles.get(font.underline, f"未知下划线样式({font.underline})")
            formats.append(f"下划线: {underline_style}")
            
            # 特别标记波浪线和点状线
            wavy_styles = [style for style in [
                getattr(WD_UNDERLINE, 'WAVY', None),
                getattr(WD_UNDERLINE, 'WAVY_HEAVY', None),
                getattr(WD_UNDERLINE, 'WAVY_DOUBLE', None)
            ] if style is not None]
            
            dotted_styles = [style for style in [
                getattr(WD_UNDERLINE, 'DOTTED', None),
                getattr(WD_UNDERLINE, 'DOTTED_HEAVY', None)
            ] if style is not None]
            
            if font.underline in wavy_styles:
                formats.append("⚠️ 波浪线格式")
            elif font.underline in dotted_styles:
                formats.append("⚠️ 点状线格式")
        
        # 上标下标
        if font.superscript:
            formats.append("上标")
        if font.subscript:
            formats.append("下标")
        
        # 删除线
        if font.strike:
            formats.append("删除线")
        
        # 粗体斜体
        if font.bold:
            formats.append("粗体")
        if font.italic:
            formats.append("斜体")
        
        # 字体颜色
        if font.color and font.color.rgb:
            try:
                rgb_val = font.color.rgb
                if hasattr(rgb_val, '__int__'):
                    color_hex = f"#{int(rgb_val):06x}"
                else:
                    color_hex = str(rgb_val)
                formats.append(f"字体颜色: {color_hex}")
            except Exception:
                formats.append("字体颜色: 特殊颜色")
        
        # 字体大小
        if font.size:
            try:
                size_pt = font.size.pt
                formats.append(f"字号: {size_pt}磅")
            except:
                formats.append("字号: 自定义大小")
        
        # 字体名称
        if font.name:
            formats.append(f"字体: {font.name}")
        
        # 检查着重号
        emphasis_mark = self._check_emphasis_mark(run)
        if emphasis_mark:
            formats.append(f"着重号: {emphasis_mark}")
        
        # 统计格式使用
        for fmt in formats:
            self.format_statistics[fmt] += 1
        
        # 保存特殊格式的文本
        if formats and run.text.strip():
            self.special_formatted_text.append({
                'text': run.text,
                'paragraph': para_index,
                'run': run_index,
                'formats': formats
            })
        
        return formats
    
    def _check_emphasis_mark(self, run):
        """检查着重号（加点字）"""
        try:
            run_xml = run._element
            em_elements = run_xml.xpath('.//w:em', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            if em_elements:
                em_val = em_elements[0].get(qn('w:val'))
                emphasis_types = {
                    'dot': '点',
                    'comma': '逗号',
                    'circle': '圆圈',
                    'underDot': '下点'
                }
                return emphasis_types.get(em_val, f'着重号({em_val})')
        except:
            pass
        
        return None
    
    def extract_format_analysis(self, docx_path):
        """提取文档的格式分析信息"""
        if not self.format_detection_enabled:
            print("🚫 格式检测功能未启用，跳过格式分析")
            return None
        
        print("🔍 开始分析文档格式...")
        
        try:
            doc = Document(docx_path)
            paragraph_count = 0
            
            # 重置统计信息
            self.special_formatted_text = []
            self.format_statistics = defaultdict(int)
            
            for para in doc.paragraphs:
                paragraph_count += 1
                for run_index, run in enumerate(para.runs):
                    self._analyze_text_formatting(run, paragraph_count, run_index)
            
            # 分析表格
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        for para in cell.paragraphs:
                            for run_index, run in enumerate(para.runs):
                                location = f"表格{table_count}-行{row_index}-列{cell_index}"
                                self._analyze_text_formatting(run, location, run_index)
            
            print(f"✅ 格式分析完成: {paragraph_count}个段落, {table_count}个表格")
            print(f"📊 发现 {len(self.special_formatted_text)} 个包含特殊格式的文本片段")
            
            return {
                'total_paragraphs': paragraph_count,
                'total_tables': table_count,
                'special_format_count': len(self.special_formatted_text),
                'format_statistics': dict(self.format_statistics)
            }
            
        except Exception as e:
            print(f"❌ 格式分析失败: {e}")
            return None
    
    def _save_format_analysis(self, format_analysis, file_path):
        """保存格式分析结果"""
        try:
            # 创建格式分析结果目录
            format_dir = Path("format_analysis")
            format_dir.mkdir(exist_ok=True)
            
            doc_name = Path(file_path).stem
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # 保存格式统计摘要
            summary_file = format_dir / f"format_summary_{doc_name}_{timestamp}.txt"
            with open(summary_file, 'w', encoding='utf-8') as f:
                f.write(f"文档格式分析报告\n")
                f.write(f"文档: {file_path}\n")
                f.write(f"分析时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"=" * 50 + "\n\n")
                
                f.write("格式统计:\n")
                for fmt, count in format_analysis['format_statistics'].items():
                    f.write(f"  {fmt}: {count}次\n")
                
                f.write(f"\n特殊格式文本详情:\n")
                for item in self.special_formatted_text:
                    f.write(f"\n位置{item['paragraph']}-片段{item['run']}: \"{item['text'][:100]}{'...' if len(item['text']) > 100 else ''}\"\n")
                    for fmt in item['formats']:
                        f.write(f"  └─ {fmt}\n")
            
            print(f"📋 格式分析报告已保存: {summary_file}")
            
            # 保存JSON格式的详细数据
            json_file = format_dir / f"format_details_{doc_name}_{timestamp}.json"
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'document_path': file_path,
                    'analysis_summary': format_analysis,
                    'special_formatted_text': self.special_formatted_text,
                    'analysis_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                }, f, ensure_ascii=False, indent=2)
            
            print(f"📄 详细格式数据已保存: {json_file}")
            
        except Exception as e:
            print(f"⚠️ 保存格式分析结果失败: {e}")
    
    def _integrate_format_info(self, api_result, format_analysis, file_path):
        """将格式信息整合到API结果中"""
        try:
            # 为API结果添加格式分析信息
            if isinstance(api_result, list):
                # 为每个题目添加格式信息概要
                format_summary = {
                    'total_special_formats': len(self.special_formatted_text),
                    'format_types': list(format_analysis['format_statistics'].keys()),
                    'analysis_enabled': True
                }
                
                # 尝试匹配特殊格式文本到具体题目
                for question in api_result:
                    question['format_info'] = format_summary.copy()
                    question['special_formats_in_question'] = []
                    
                    # 检查题目内容是否包含特殊格式的文本
                    question_text = question.get('question', {}).get('content', '')
                    for format_item in self.special_formatted_text:
                        if format_item['text'].strip() in question_text:
                            question['special_formats_in_question'].append({
                                'text': format_item['text'],
                                'formats': format_item['formats']
                            })
            
            print(f"🔗 格式信息已整合到API结果中")
            
        except Exception as e:
            print(f"⚠️ 整合格式信息失败: {e}")
    
    def get_special_format_summary(self):
        """获取特殊格式摘要信息"""
        if not self.format_detection_enabled:
            return "格式检测功能未启用"
        
        if not self.special_formatted_text:
            return "未发现特殊格式文本"
        
        summary = []
        summary.append(f"📊 特殊格式统计 (总计: {len(self.special_formatted_text)} 个片段):")
        
        # 按格式类型统计
        format_counts = defaultdict(int)
        for item in self.special_formatted_text:
            for fmt in item['formats']:
                if '⚠️' in fmt:  # 重点关注的格式
                    format_counts[fmt] += 1
        
        for fmt, count in sorted(format_counts.items(), key=lambda x: x[1], reverse=True):
            summary.append(f"  {fmt}: {count}次")
        
        return "\n".join(summary)
    
    def get_supported_formats(self):
        """获取支持的文档格式"""
        return ['.docx', '.doc', '.rtf', '.odt', '.txt']
    
    def is_supported_format(self, file_path):
        """检查文件格式是否支持"""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.get_supported_formats()
    
    def extract_images_from_docx(self, docx_path, save_images=False):
        """从docx文件中提取图片"""
        print(f"🖼️  从docx文件中提取图片...")
        
        images = []
        if save_images:
            # 创建media文件夹
            media_dir = Path("media")
            media_dir.mkdir(exist_ok=True)
            print(f"📁 创建图片保存目录: {media_dir}")
        
        try:
            # docx文件本质上是一个zip文件
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                # 查找media文件夹中的图片
                for file_info in zip_file.filelist:
                    if file_info.filename.startswith('word/media/'):
                        file_name = file_info.filename.split('/')[-1]
                        file_ext = Path(file_name).suffix.lower()
                        
                        # 只处理图片文件
                        if file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                            try:
                                # 读取图片数据
                                with zip_file.open(file_info.filename) as img_file:
                                    img_data = img_file.read()
                                
                                # 转换为PIL Image对象
                                img = Image.open(BytesIO(img_data))
                                
                                # 如果需要保存图片到本地
                                if save_images:
                                    img_path = media_dir / file_name
                                    with open(img_path, 'wb') as f:
                                        f.write(img_data)
                                    print(f"  💾 保存图片: {img_path}")
                                
                                images.append({
                                    'filename': file_name,
                                    'path': file_info.filename,
                                    'image': img,
                                    'data': img_data,
                                    'size': img.size,
                                    'format': img.format
                                })
                                
                                print(f"  📷 提取图片: {file_name} ({img.size[0]}x{img.size[1]})")
                                
                            except Exception as e:
                                print(f"  ⚠️  图片 {file_name} 读取失败: {e}")
                                continue
                
                print(f"✅ 共提取 {len(images)} 张图片")
                return images
                
        except Exception as e:
            print(f"❌ 提取图片失败: {e}")
            return []
    
    def analyze_image_with_llm(self, image_data, image_name):
        """使用LLM分析图片内容"""
        print(f"  🤖 使用LLM分析图片: {image_name}")
        
        try:
            # 将图片转换为base64
            img_buffer = BytesIO()
            if isinstance(image_data, bytes):
                # 如果已经是bytes，直接使用
                img_base64 = base64.b64encode(image_data).decode()
            else:
                # 如果是PIL Image，先保存为bytes
                image_data.save(img_buffer, format='PNG')
                img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
            
            # 构建API请求
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            # 使用vision API分析图片
            data = {
                "model": "doubao-seed-1-6-250615",
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "请分析这张图片的内容，用简洁的中文描述图片中显示的内容。如果是试卷题目，请描述题目类型和主要内容。"
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_base64}"
                                }
                            }
                        ]
                    }
                ],
                "max_tokens": 500,
                "temperature": 0.1
            }
            
            # 调用API
            response = requests.post(
                f"{self.base_url}/v3/chat/completions",
                headers=headers,
                json=data,
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    content = result['choices'][0]['message']['content']
                    print(f"  ✅ 图片分析完成: {content[:100]}...")
                    return content
                else:
                    print(f"  ⚠️  API响应格式异常")
                    return f"图片内容: {image_name}"
            else:
                print(f"  ❌ API调用失败: {response.status_code}")
                return f"图片内容: {image_name}"
                
        except Exception as e:
            print(f"  ❌ 图片分析异常: {e}")
            return f"图片内容: {image_name}"
    
    def replace_image_watermarks(self, content, images):
        """替换内容中的图片水印为图片内容描述"""
        print("🔄 替换图片水印...")
        
        if not images:
            print("  ℹ️  没有图片需要处理")
            return content
        
        # 查找图片引用模式
        # 匹配类似 ![学科网(www.zxxk.com)--教育资源门户...](media/image6.png) 的模式
        image_pattern = r'!\[([^\]]+)\]\(([^)]+)\)'
        
        def replace_image_ref(match):
            watermark_text = match.group(1)
            image_path = match.group(2)
            
            # 提取图片文件名
            image_filename = Path(image_path).name
            
            # 查找对应的图片
            for img_info in images:
                if img_info['filename'] == image_filename:
                    # 使用LLM分析图片内容
                    image_description = self.analyze_image_with_llm(img_info['image'], image_filename)
                    
                    # 替换水印文字为图片描述
                    new_text = f"![{image_description}]({image_path})"
                    print(f"  🔄 替换: {watermark_text[:50]}... -> {image_description[:50]}...")
                    return new_text
            
            # 如果没找到对应图片，保留原样
            print(f"  ⚠️  未找到图片: {image_filename}")
            return match.group(0)
        
        # 执行替换
        modified_content = re.sub(image_pattern, replace_image_ref, content)
        
        # 统计替换次数
        original_count = len(re.findall(image_pattern, content))
        modified_count = len(re.findall(image_pattern, modified_content))
        
        print(f"✅ 水印替换完成，处理了 {len(images)} 张图片")
        return modified_content
    
    def convert_word_to_text(self, file_path, output_format='markdown'):
        """使用pandoc将Word文档转换为文本，并增强格式标注"""
        if not self.pandoc_available:
            print("❌ Pandoc不可用，无法处理文档")
            return None
        
        if not self.is_supported_format(file_path):
            print(f"❌ 不支持的文件格式: {Path(file_path).suffix}")
            print(f"支持的格式: {', '.join(self.get_supported_formats())}")
            return None
        
        if not os.path.exists(file_path):
            print(f"❌ 文件不存在: {file_path}")
            return None
        
        print(f"📄 开始处理文档: {file_path}")
        print(f"📊 文件大小: {os.path.getsize(file_path) / (1024*1024):.2f} MB")
        
        try:
            # 构建pandoc命令
            cmd = [
                'pandoc',
                file_path,
                '--to', output_format,
                '--wrap', 'none',  # 不自动换行
                '--standalone',  # 生成独立文档
                '--quiet'  # 减少输出
            ]
            
            print(f"🔧 执行命令: {' '.join(cmd)}")
            
            # 执行pandoc转换
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            
            if result.returncode == 0:
                content = result.stdout
                print(f"✅ 转换成功: {len(content)} 字符")
                
                # 如果是docx文件，提取图片并替换水印
                if file_path.lower().endswith('.docx'):
                     print("🖼️  检测到docx文件，开始处理图片...")
                     images = self.extract_images_from_docx(file_path, save_images=True)
                     if images:
                         content = self.replace_image_watermarks(content, images)
                     else:
                         print("ℹ️  未找到图片或图片处理失败")
                
                # 新增：如果有格式分析结果，增强pandoc内容
                if hasattr(self, 'special_formatted_text') and self.special_formatted_text:
                    print("🎨 开始增强格式标注...")
                    content = self._enhance_content_with_format_info(content)
                
                # 保存转换结果
                pandoc_res_dir = Path("pandoc_res")
                pandoc_res_dir.mkdir(exist_ok=True)
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = pandoc_res_dir / f"pandoc转换结果_{timestamp}.txt"
                with open(output_filename, 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"📝 转换结果已保存到: {output_filename}")
                
                return content
            else:
                print(f"❌ 转换失败: {result.stderr}")
                return None
                
        except subprocess.TimeoutExpired:
            print("❌ 转换超时（5分钟）")
            return None
        except Exception as e:
            print(f"❌ 转换异常: {e}")
            return None
    
    def _enhance_content_with_format_info(self, content):
        """根据格式分析结果增强pandoc内容"""
        print(f"🔍 开始分析 {len(self.special_formatted_text)} 个特殊格式文本")
        
        # 按文本长度排序，从长到短，避免短文本替换影响长文本
        sorted_formats = sorted(self.special_formatted_text, 
                              key=lambda x: len(x['text']), reverse=True)
        
        enhanced_count = 0
        
        for format_item in sorted_formats:
            text = format_item['text'].strip()
            formats = format_item['formats']
            
            # 跳过空文本或过短文本
            if len(text) < 2:
                continue
            
            # 生成格式标注
            format_annotation = self._generate_format_annotation(formats)
            
            if format_annotation and text in content:
                # 创建增强的文本标注
                enhanced_text = f"[{text}]{{{format_annotation}}}"
                
                # 执行替换（只替换第一个匹配，避免重复替换）
                content = content.replace(text, enhanced_text, 1)
                enhanced_count += 1
                
                print(f"  ✨ 增强: \"{text[:30]}{'...' if len(text) > 30 else ''}\" -> {format_annotation}")
        
        print(f"✅ 格式增强完成，共处理 {enhanced_count} 个文本")
        return content
    
    def _generate_format_annotation(self, formats):
        """根据格式列表生成标注"""
        annotations = []
        
        # 处理下划线类型
        for fmt in formats:
            if "波浪线格式" in fmt:
                annotations.append(".wavy-underline")
            elif "点状线格式" in fmt:
                annotations.append(".dotted-underline")
            elif "下划线: 单下划线" in fmt:
                annotations.append(".single-underline")
            elif "下划线: 双下划线" in fmt:
                annotations.append(".double-underline")
            elif "下划线: 粗下划线" in fmt:
                annotations.append(".thick-underline")
            elif "下划线: 虚线下划线" in fmt:
                annotations.append(".dashed-underline")
            elif "删除线" in fmt:
                annotations.append(".strikethrough")
            elif "上标" in fmt:
                annotations.append(".superscript")
            elif "下标" in fmt:
                annotations.append(".subscript")
            elif "着重号" in fmt:
                annotations.append(".emphasis-mark")
            elif "粗体" in fmt:
                annotations.append(".bold")
            elif "斜体" in fmt:
                annotations.append(".italic")
        
        # 处理字体颜色（提取颜色值）
        for fmt in formats:
            if "字体颜色:" in fmt and "000000" not in fmt:  # 跳过黑色
                color = fmt.split("字体颜色:")[-1].strip()
                annotations.append(f".color-{color}")
        
        # 处理字号
        for fmt in formats:
            if "字号:" in fmt and "磅" in fmt:
                size = fmt.split("字号:")[-1].replace("磅", "").strip()
                try:
                    size_num = float(size)
                    if size_num != 12.0:  # 跳过默认字号
                        annotations.append(f".font-{size}pt")
                except:
                    pass
        
        return " ".join(annotations) if annotations else None
    
    def convert_word_to_markdown(self, file_path):
        """将Word文档转换为Markdown格式"""
        return self.convert_word_to_text(file_path, 'markdown')
    
    def convert_word_to_plain_text(self, file_path):
        """将Word文档转换为纯文本格式"""
        return self.convert_word_to_text(file_path, 'plain')
    
    def convert_word_to_html(self, file_path):
        """将Word文档转换为HTML格式"""
        return self.convert_word_to_text(file_path, 'html')
    
    def call_llm_api(self, content, prompt_template_path="prompt.md"):
        """调用大模型API解析文档结构"""
        print("🤖 开始调用大模型API...")
        
        # 读取prompt模板
        try:
            with open(prompt_template_path, 'r', encoding='utf-8') as f:
                prompt_template = f.read()
            # 使用安全的字符串替换
            prompt = prompt_template.replace("{content}", content)
            print(f"✅ 成功加载prompt模板: {prompt_template_path}")
        except FileNotFoundError:
            print(f"⚠️ 未找到prompt模板文件: {prompt_template_path}")
            print("使用默认prompt...")
            prompt = f"""
一套试卷有三级结构，1. 分题型/类型的大模块 2.完整的一道题 3. 完整的一道题中的多个小题。你需要解析后两级结构。
请分析以下文档内容，提取出试卷的二级结构（完整的一道题），返回JSON格式的数组，每个对象包含以下字段：
- total_number: 总题号，唯一，及此题按照试卷题目展示顺序的总题号。（字符串）
- module_number: 模块中的题号，即在一级结构中的题号。（字符串） 
- question_first_sentence: 题目的第一个分句（字符串）
- question_page: 题目所在的页码（int数组）
- answer_first_sentence: 题目答案的第一个分句（字符串，如果没有答案则为空字符串）
- explanation_first_sentence: 题目解析的第一个分句（字符串，如果没有解析则为空字符串）
- answer_page: 题目答案解析所在的页码（int数组，如果没有答案则为null）

文档内容：
{content}

请只返回JSON数组，不要包含其他说明文字。
如果题目/答案跨页，则在页码中需要记录两页信息。
"""
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        # 先测试API连接
        test_data = {
            "model": "doubao-seed-1-6-250615",
            "messages": [{"role": "user", "content": "Hello"}],
            "max_tokens": 32000,
            "temperature": 0.1,
            "stream": True,
            "thinking": {
                "type": "enabled",
                "budget_tokens": 2000
            }
        }
        
        print("🔍 测试API连接...")
        try:
            test_response = requests.post(
                f"{self.base_url}/v3/chat/completions", 
                headers=headers, 
                json=test_data, 
                timeout=30
            )
            test_response.raise_for_status()
            print("✅ API连接测试成功")
        except requests.exceptions.HTTPError as e:
            print(f"❌ API连接测试失败: {e}")
            print(f"响应状态码: {test_response.status_code}")
            print(f"响应内容: {test_response.text}")
            return None
        except Exception as e:
            print(f"❌ API连接测试异常: {e}")
            return None
        
        # 正式请求 - 使用流式输出
        data = {
            "model": "doubao-seed-1-6-250615",
            "messages": [{"role": "user", "content": prompt}],
            "stream": True,
            "thinking": {
                "type": "enabled",
                "budget_tokens": 1500
            },
            "response_format": {
                "type": "json_object"
            },
            "temperature": 0.1,
            "max_completion_tokens": 32000
        }
        
        print("🚀 调用大模型API解析文档结构...")
        print(f"请求URL: {self.base_url}/v3/chat/completions")
        print(f"模型: {data['model']}")
        print(f"消息长度: {len(prompt)} 字符")
        
        try:
            response = requests.post(
                f"{self.base_url}/v3/chat/completions", 
                headers=headers, 
                json=data, 
                timeout=300,  # 减少超时时间到5分钟
                stream=True
            )
            response.raise_for_status()
            
            # 处理流式响应
            llm_content = ""
            print("📡 开始接收流式响应...")
            
            for line in response.iter_lines():
                if line:
                    line = line.decode('utf-8')
                    if line.startswith('data: '):
                        data_str = line[6:]  # 去掉 'data: ' 前缀
                        
                        if data_str == '[DONE]':
                            print("\n✅ 流式响应接收完成")
                            break
                        
                        try:
                            data_json = json.loads(data_str)
                            if 'choices' in data_json and len(data_json['choices']) > 0:
                                choice = data_json['choices'][0]
                                
                                # 处理thinking状态
                                if 'thinking' in choice:
                                    thinking = choice['thinking']
                                    if thinking.get('type') == 'thinking':
                                        print(f"🤔 思考中... ({thinking.get('tokens_used', 0)} tokens)")
                                    elif thinking.get('type') == 'finished':
                                        print(f"✅ 思考完成，共使用 {thinking.get('tokens_used', 0)} tokens")
                                
                                # 处理delta内容
                                if 'delta' in choice and 'content' in choice['delta']:
                                    content = choice['delta']['content']
                                    llm_content += content
                                    print(content, end='', flush=True)
                                    
                        except json.JSONDecodeError:
                            continue
            
            print(f"\n✅ API调用成功，响应长度: {len(llm_content)} 字符")
            return llm_content
            
        except requests.exceptions.HTTPError as e:
            print(f"HTTP错误: {e}")
            print(f"响应状态码: {response.status_code}")
            print(f"响应内容: {response.text}")
            return None
        except Exception as e:
            print(f"❌ API调用异常: {e}")
            return None
    
    def process_word_document(self, file_path, output_format='markdown', prompt_template_path="prompt.md", 
                            enable_format_analysis=True, enable_dot_below_detection=True):
        """完整的Word文档处理流程"""
        print("=" * 60)
        print("Pandoc Word文档处理工具 - 增强版 (支持加点字)")
        print("=" * 60)
        print(f"文档文件: {file_path}")
        print(f"输出格式: {output_format}")
        print(f"Prompt模板: {prompt_template_path}")
        print(f"格式分析: {'启用' if enable_format_analysis else '禁用'}")
        print(f"加点字检测: {'启用' if enable_dot_below_detection else '禁用'}")
        print("=" * 60)
        
        # 🆕 新增：第一步 - 加点字预处理（如果启用且为docx文件）
        processed_file_path = file_path
        if enable_dot_below_detection and file_path.lower().endswith('.docx'):
            processed_file_path = self._preprocess_dot_below_chars(file_path)
            if not processed_file_path:
                processed_file_path = file_path  # 回退到原文件
        
        # 第二步 - 格式分析（如果启用且为docx文件）
        format_analysis = None
        if enable_format_analysis and processed_file_path.lower().endswith('.docx'):
            format_analysis = self.extract_format_analysis(processed_file_path)
            if format_analysis:
                self._save_format_analysis(format_analysis, processed_file_path)
        
        # 第三步：使用pandoc转换文档
        content = self.convert_word_to_text(processed_file_path, output_format)
        if not content:
            print("❌ 文档转换失败，无法继续处理")
            return None
        
        # 🆕 第四步：转换加点字标记为HTML格式
        if enable_dot_below_detection:
            content = self._convert_dot_below_markers_to_html(content)
        
        # 第三步：调用大模型API解析内容
        llm_response = self.call_llm_api(content, prompt_template_path)
        if not llm_response:
            print("❌ API调用失败")
            return None
        
        # 第四步：处理API响应并集成格式信息
        api_result = self._process_api_response(llm_response, file_path)
        
        # 第五步：如果有格式分析结果，将其整合到最终结果中
        if format_analysis and api_result:
            self._integrate_format_info(api_result, format_analysis, file_path)
        
        return api_result
    
    def _process_api_response(self, llm_content, original_file_path):
        """处理API响应并保存结果"""
        # 去掉markdown代码块标记（如果有的话）
        if '```json' in llm_content:
            start_idx = llm_content.find('```json') + 7
            end_idx = llm_content.find('```', start_idx)
            if end_idx != -1:
                llm_content = llm_content[start_idx:end_idx]
            else:
                llm_content = llm_content[start_idx:]
        elif '```' in llm_content:
            parts = llm_content.split('```')
            if len(parts) >= 2:
                llm_content = parts[1]
        
        llm_content = llm_content.strip()
        
        # 检测重复模式（防止API生成错误）
        import re
        # 检测重复的score模式
        if re.search(r'("score"\s*,\s*){3,}', llm_content):
            print("❌ 检测到重复的score模式，API响应异常")
            return None
        
        # 检测其他重复模式
        if re.search(r'(\s*,\s*,\s*){3,}', llm_content):
            print("❌ 检测到重复的逗号模式，API响应异常")
            return None
        
        # 修复JSON中的双引号问题
        # 替换题目内容中的未转义双引号
        llm_content = re.sub(r'将数据"([^"]+)"', r'将数据\\"\\1\\"', llm_content)
        
        # 保存原始API响应
        # timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # raw_filename = f"raw_api_response_{timestamp}.txt"
        # with open(raw_filename, 'w', encoding='utf-8') as f:
        #     f.write(llm_content)
        # print(f"📄 原始API响应已保存到: {raw_filename}")
        
        # 尝试转换为JSON
        try:
            questions = json.loads(llm_content)
            
            # 创建 json_res 文件夹
            json_res_dir = Path("json_res")
            json_res_dir.mkdir(exist_ok=True)
            
            output_file = json_res_dir / f"questions_with_pandoc_{Path(original_file_path).stem}.json"
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(questions, f, ensure_ascii=False, indent=2)
            print(f"🎉 完成！共{len(questions)}道题目，保存到: {output_file}")
            return questions
        except json.JSONDecodeError as e:
            print(f"❌ JSON解析失败: {e}")
            print(f"错误位置: 第{e.lineno}行第{e.colno}列")
            print("请检查原始API响应文件")
            return None
    
    def _preprocess_dot_below_chars(self, docx_path):
        """预处理docx文件中的加点字，使pandoc能识别"""
        print("🔍 预处理加点字...")
        
        try:
            # 导入预处理器
            import zipfile
            import xml.etree.ElementTree as ET
            import tempfile
            import shutil
            import re
            
            output_path = docx_path.replace('.docx', '_dot_processed.docx')
            
            # 创建临时目录来解压和重新打包docx
            with tempfile.TemporaryDirectory() as temp_dir:
                extract_dir = Path(temp_dir) / "docx_content"
                extract_dir.mkdir()
                
                # 解压docx文件
                with zipfile.ZipFile(docx_path, 'r') as zip_file:
                    zip_file.extractall(extract_dir)
                
                # 修改document.xml
                document_xml_path = extract_dir / "word" / "document.xml"
                if document_xml_path.exists():
                    with open(document_xml_path, 'r', encoding='utf-8') as f:
                        xml_content = f.read()
                    
                    # 查找并替换加点字标记
                    run_with_em_pattern = r'(<w:r>.*?<w:rPr>.*?)<w:em w:val="dot"\s*/>(.*?</w:rPr>.*?<w:t>)(.*?)(</w:t>.*?</w:r>)'
                    
                    def replace_run_with_em(match):
                        before_em = match.group(1)
                        after_em = match.group(2) 
                        text_content = match.group(3)
                        after_text = match.group(4)
                        
                        # 添加下划线和特殊标记
                        underline_xml = '<w:u w:val="single"/>'
                        marked_text = f"[DOT_BELOW]{text_content}[/DOT_BELOW]"
                        
                        return f"{before_em}{underline_xml}{after_em}{marked_text}{after_text}"
                    
                    modified_content, replacement_count = re.subn(run_with_em_pattern, replace_run_with_em, xml_content, flags=re.DOTALL)
                    
                    if replacement_count > 0:
                        with open(document_xml_path, 'w', encoding='utf-8') as f:
                            f.write(modified_content)
                        print(f"  ✅ 处理了 {replacement_count} 个加点字")
                
                # 重新打包docx文件
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for file_path in extract_dir.rglob('*'):
                        if file_path.is_file():
                            relative_path = file_path.relative_to(extract_dir)
                            zip_file.write(file_path, relative_path)
            
            return output_path
            
        except Exception as e:
            print(f"  ⚠️ 预处理失败: {e}")
            return None
    
    def _convert_dot_below_markers_to_html(self, content):
        """将加点字标记转换为HTML格式"""
        print("🎨 转换加点字标记为HTML格式...")
        
        try:
            import re
            
            # 匹配模式：[\[DOT_BELOW\]字符\[/DOT_BELOW\]]{.underline}
            pattern = r'\[\\\[DOT_BELOW\\\]([\u4e00-\u9fff])\\\[/DOT_BELOW\\\]\]\{\.underline\}'
            
            def replace_with_html(match):
                char = match.group(1)
                return f'<span style="text-emphasis: filled dot black; text-emphasis-position: under right;" data-mce-style="text-emphasis: filled dot black; text-emphasis-position: under right;">{char}</span>'
            
            converted_content, count = re.subn(pattern, replace_with_html, content)
            
            if count > 0:
                print(f"  ✅ 转换了 {count} 个加点字为HTML格式")
            
            return converted_content
            
        except Exception as e:
            print(f"  ⚠️ 加点字转换失败: {e}")
            return content

def main():
    """主函数"""
    import sys
    
    # 配置参数
    if len(sys.argv) > 1:
        word_file_path = sys.argv[1]
    else:
        word_file_path = "Chinese/精品解析：2025年甘肃省兰州市中考语文真题（解析版）.docx"  # 默认文件路径
     
    output_format = "markdown"  # 可选: markdown, plain, html
    prompt_template_path = "prompt_Chinese.md"
    
    # 创建处理器实例
    processor = PandocWordProcessor()
    
    # 检查pandoc可用性
    if not processor.pandoc_available:
        print("❌ Pandoc不可用，请先安装pandoc")
        print("安装方法:")
        print("  macOS: brew install pandoc")
        print("  Ubuntu/Debian: sudo apt-get install pandoc")
        print("  Windows: 下载安装包 https://pandoc.org/installing.html")
        return
    
    # 检查文件格式
    if not processor.is_supported_format(word_file_path):
        print(f"❌ 不支持的文件格式: {Path(word_file_path).suffix}")
        print(f"请使用以下格式之一: {', '.join(processor.get_supported_formats())}")
        return
    
    # 处理文档
    result = processor.process_word_document(
        word_file_path, 
        output_format, 
        prompt_template_path
    )
    
    if result:
        print("✅ 文档处理完成！")
        
        # 显示特殊格式摘要
        format_summary = processor.get_special_format_summary()
        print("\n" + "="*50)
        print(format_summary)
        print("="*50)
    else:
        print("❌ 文档处理失败")

if __name__ == "__main__":
    main() 