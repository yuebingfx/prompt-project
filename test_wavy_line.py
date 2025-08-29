#!/usr/bin/env python3
"""
波浪线下划线检测测试脚本
测试原始文档中的波浪线格式识别和处理情况
"""

def test_original_document():
    """测试原始Word文档中的波浪线格式"""
    print("=" * 60)
    print("🔍 测试原始Word文档中的波浪线格式")
    print("=" * 60)
    
    try:
        from docx import Document
        
        # 读取原始文档
        doc_path = "Chinese/精品解析：2025年四川省宜宾市中考语文真题（解析版）.docx"
        doc = Document(doc_path)
        print(f"📄 正在分析文档: {doc_path}")
        
        found_wavy = False
        wavy_runs = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if '然而' in text and '勾勒' in text:
                print(f"\n📍 找到目标段落 {i+1}: {repr(text[:50])}...")
                
                for j, run in enumerate(para.runs):
                    run_text = run.text
                    font = run.font
                    
                    print(f"  Run {j}: {repr(run_text)} (长度: {len(run_text)})")
                    
                    if font.underline:
                        print(f"    ✅ 下划线格式: {font.underline}")
                        
                        # 检查是否是波浪线格式
                        try:
                            from docx.enum.text import WD_UNDERLINE
                            wavy_styles = [
                                getattr(WD_UNDERLINE, 'WAVY', None),
                                getattr(WD_UNDERLINE, 'WAVY_HEAVY', None), 
                                getattr(WD_UNDERLINE, 'WAVY_DOUBLE', None)
                            ]
                            wavy_styles = [s for s in wavy_styles if s is not None]
                            
                            if font.underline in wavy_styles:
                                print(f"    🌊 发现波浪线下划线! 样式值: {font.underline}")
                                found_wavy = True
                                wavy_runs.append({
                                    'para': i+1,
                                    'run': j,
                                    'text': repr(run_text),
                                    'length': len(run_text),
                                    'style': font.underline
                                })
                            else:
                                print(f"    ➡️ 其他下划线格式: {font.underline}")
                        except Exception as e:
                            print(f"    ⚠️ 格式检查失败: {e}")
                    else:
                        print(f"    ❌ 无下划线格式")
                
                break
        
        print(f"\n📊 波浪线检测结果:")
        if found_wavy:
            print(f"✅ 发现 {len(wavy_runs)} 个波浪线下划线段落")
            for item in wavy_runs:
                print(f"  - 段落{item['para']}, Run{item['run']}: {item['text']} (长度{item['length']}, 样式{item['style']})")
        else:
            print("❌ 未发现波浪线下划线")
            
    except Exception as e:
        print(f"❌ 测试失败: {e}")

def test_processed_document():
    """测试处理后的Word文档"""
    print("\n" + "=" * 60)
    print("🔍 测试处理后的Word文档")
    print("=" * 60)
    
    try:
        from docx import Document
        import os
        
        # 检查processed文件夹中的文档
        processed_path = "Chinese/processed/精品解析：2025年四川省宜宾市中考语文真题（解析版）_dot_processed.docx"
        
        if not os.path.exists(processed_path):
            print(f"❌ 处理后的文档不存在: {processed_path}")
            return
            
        doc = Document(processed_path)
        print(f"📄 正在分析处理后文档: {processed_path}")
        
        found_wavy_chars = False
        wavy_char_runs = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if '然而' in text and ('勾勒' in text or '～' in text or '~' in text):
                print(f"\n📍 找到目标段落 {i+1}: {repr(text[:80])}...")
                
                for j, run in enumerate(para.runs):
                    run_text = run.text
                    
                    if '～' in run_text or '~' in run_text:
                        print(f"  🌊 Run {j}: {repr(run_text)} - 发现波浪线字符!")
                        found_wavy_chars = True
                        wavy_char_runs.append({
                            'para': i+1,
                            'run': j,
                            'text': repr(run_text),
                            'wavy_count': run_text.count('～') + run_text.count('~')
                        })
                    else:
                        print(f"  ➡️ Run {j}: {repr(run_text)}")
                
                break
        
        print(f"\n📊 波浪线字符检测结果:")
        if found_wavy_chars:
            print(f"✅ 发现 {len(wavy_char_runs)} 个包含波浪线字符的Run")
            for item in wavy_char_runs:
                print(f"  - 段落{item['para']}, Run{item['run']}: {item['text']} (波浪线数量: {item['wavy_count']})")
        else:
            print("❌ 未发现波浪线字符")
            
    except Exception as e:
        print(f"❌ 测试失败: {e}")

def test_pandoc_output():
    """测试pandoc转换后的输出"""
    print("\n" + "=" * 60)
    print("🔍 测试Pandoc转换后的输出")
    print("=" * 60)
    
    try:
        import os
        import glob
        
        # 查找最新的pandoc转换结果
        pandoc_files = glob.glob("pandoc_res/pandoc转换结果_*.txt")
        if not pandoc_files:
            print("❌ 未找到pandoc转换结果文件")
            return
            
        latest_file = max(pandoc_files, key=os.path.getmtime)
        print(f"📄 正在分析最新的pandoc转换结果: {latest_file}")
        
        with open(latest_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 查找目标段落
        lines = content.split('\n')
        target_lines = []
        
        for i, line in enumerate(lines):
            if '然而' in line and '勾勒' in line:
                target_lines.append((i+1, line))
                print(f"\n📍 找到目标行 {i+1}: {repr(line)}")
                
                # 分析这一行的内容
                if '～' in line or '~' in line:
                    wavy_count = line.count('～') + line.count('~')
                    print(f"  🌊 发现波浪线字符: {wavy_count} 个")
                else:
                    # 检查空格模式
                    import re
                    space_matches = re.findall(r'\s{3,}', line)
                    if space_matches:
                        print(f"  📏 发现长空格序列: {len(space_matches)} 个")
                        for j, match in enumerate(space_matches):
                            print(f"    - 空格序列{j+1}: 长度{len(match)}")
                    else:
                        print(f"  ❌ 未发现波浪线字符或长空格")
        
        print(f"\n📊 Pandoc输出检测结果:")
        if target_lines:
            print(f"✅ 找到 {len(target_lines)} 个目标行")
        else:
            print("❌ 未找到包含'然而'和'勾勒'的行")
            
    except Exception as e:
        print(f"❌ 测试失败: {e}")

def test_wavy_html_conversion():
    """测试波浪线HTML转换逻辑"""
    print("\n" + "=" * 60)
    print("🧪 测试波浪线HTML转换逻辑")
    print("=" * 60)
    
    try:
        from docx import Document
        from docx.enum.text import WD_UNDERLINE
        
        # 读取原始文档
        doc_path = "Chinese/精品解析：2025年四川省宜宾市中考语文真题（解析版）.docx"
        doc = Document(doc_path)
        
        print("🔍 查找波浪线下划线并模拟转换...")
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if '然而' in text and '勾勒' in text:
                print(f"\n📍 找到目标段落 {i+1}")
                
                for j, run in enumerate(para.runs):
                    run_text = run.text
                    font = run.font
                    
                    if font.underline:
                        # 检查是否是波浪线格式
                        wavy_styles = [
                            getattr(WD_UNDERLINE, 'WAVY', None),
                            getattr(WD_UNDERLINE, 'WAVY_HEAVY', None),
                            getattr(WD_UNDERLINE, 'WAVY_DOUBLE', None)
                        ]
                        wavy_styles = [s for s in wavy_styles if s is not None]
                        
                        if font.underline in wavy_styles and run_text.strip() == "" and len(run_text) > 2:
                            print(f"\n🌊 发现波浪线下划线空格序列:")
                            print(f"   原始: {repr(run_text)} (长度: {len(run_text)})")
                            
                            # 测试不同的转换方案
                            
                            # 方案1: 原始HTML方案
                            nbsp_count = max(3, len(run_text) // 3)
                            nbsp_html = "&nbsp; " * nbsp_count
                            wavy_html = f'<span style="text-decoration: underline wavy; text-underline-position: under;" data-mce-style="text-decoration: underline wavy; text-underline-position: under;">{nbsp_html.strip()}</span>'
                            
                            # 方案2: Pandoc原生格式
                            nbsp_pandoc = "&nbsp;" * nbsp_count
                            pandoc_format = f'[{nbsp_pandoc}]{{.single-wavyunderline}}'
                            
                            print(f"   转换方案对比:")
                            print(f"   方案1 HTML: {wavy_html}")
                            print(f"   方案2 Pandoc: {pandoc_format}")
                            print(f"   空格数量: {nbsp_count}")
                            
                            # 验证格式
                            print("\n   🔧 测试两种方案的Pandoc处理效果:")
                            print("   方案1 (HTML标签):")
                            test_pandoc_html(wavy_html)
                            
                            print("   方案2 (Pandoc原生):")
                            test_pandoc_html(pandoc_format)
                            
                            return [wavy_html, pandoc_format]
                break
        
        print("❌ 未找到波浪线下划线")
        return None
        
    except Exception as e:
        print(f"❌ 测试失败: {e}")
        return None

def test_pandoc_html(html_content):
    """测试pandoc如何处理我们的HTML"""
    try:
        import subprocess
        import tempfile
        import os
        
        # 创建临时HTML文件
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
            test_html = f"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8"></head>
<body>
<p>然而，{html_content}。在某种意义上，努力与付出会一笔一画勾勒出未来的轮廓。</p>
</body>
</html>"""
            f.write(test_html)
            temp_html_path = f.name
        
        try:
            # 使用pandoc转换为markdown
            result = subprocess.run([
                'pandoc', temp_html_path, '--to', 'markdown', '--wrap', 'none'
            ], capture_output=True, text=True, encoding='utf-8')
            
            if result.returncode == 0:
                output = result.stdout.strip()
                print(f"   Pandoc输出: {repr(output)}")
                
                if 'underline wavy' in output:
                    print("   ✅ Pandoc保留了波浪线样式")
                elif '&nbsp;' in output:
                    print("   ✅ Pandoc保留了非断空格")
                elif html_content in output:
                    print("   ✅ Pandoc保留了原始HTML")
                else:
                    print("   ⚠️ Pandoc可能简化了HTML")
            else:
                print(f"   ❌ Pandoc转换失败: {result.stderr}")
                
        finally:
            # 清理临时文件
            os.unlink(temp_html_path)
            
    except Exception as e:
        print(f"   ⚠️ Pandoc测试失败: {e}")

def test_different_html_approaches():
    """测试不同的HTML方案"""
    print("\n" + "=" * 60)
    print("🎯 测试不同的HTML波浪线方案")
    print("=" * 60)
    
    # 方案1: 用户提供的方案
    approach1 = '<span style="text-decoration: underline wavy; text-underline-position: under;" data-mce-style="text-decoration: underline wavy; text-underline-position: under;">&nbsp; &nbsp; &nbsp; &nbsp; &nbsp;</span>'
    
    # 方案2: 简化版本
    approach2 = '<span style="text-decoration: underline wavy;">&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</span>'
    
    # 方案3: 使用波浪线字符
    approach3 = '～～～～～'
    
    # 方案4: 使用input标签
    approach4 = '<input style="border: none; border-bottom: 2px wavy #000; background: transparent;" size="10" readonly/>'
    
    # 方案5: 用户新提供的Pandoc原生格式
    approach5 = '[&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;]{.single-wavyunderline}'
    
    approaches = [
        ("用户推荐方案", approach1),
        ("简化CSS方案", approach2), 
        ("波浪线字符方案", approach3),
        ("Input标签方案", approach4),
        ("Pandoc原生格式", approach5)
    ]
    
    for name, html in approaches:
        print(f"\n📋 {name}:")
        print(f"   HTML: {html}")
        print(f"   长度: {len(html)}")
        
        # 测试pandoc处理
        test_pandoc_html(html)

def main():
    """主测试函数"""
    print("🧪 波浪线下划线检测和处理测试")
    print("测试目标: '然而，___。在某种意义上，努力与付出会一笔一画勾勒出未来的轮廓'")
    
    # 测试1: 原始文档中的波浪线格式
    test_original_document()
    
    # 测试2: 处理后文档中的波浪线字符
    test_processed_document()
    
    # 测试3: pandoc转换后的输出
    test_pandoc_output()
    
    # 测试4: 波浪线HTML转换逻辑
    test_wavy_html_conversion()
    
    # 测试5: 不同HTML方案对比
    test_different_html_approaches()
    
    print("\n🏁 测试完成!")

if __name__ == "__main__":
    main()
