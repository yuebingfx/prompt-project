import re
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

def extract_text_with_format(doc):
    """
    从Word文档中提取文本，保留段落结构，并处理数学公式为LaTeX格式
    :param doc: python-docx的Document对象
    :return: 处理后的Markdown文本（含LaTeX公式）
    """
    markdown_content = []
    current_heading_level = 0  # 记录当前标题层级

    for paragraph in doc.paragraphs:
        para_text = paragraph.text.strip()
        if not para_text:
            continue  # 跳过空段落

        # 1. 识别标题（基于Word文档中的标题样式，如"标题1"、"标题2"）
        para_style = paragraph.style.name
        if "标题1" in para_style:
            markdown_content.append(f"# {para_text}\n")
            current_heading_level = 1
        elif "标题2" in para_style:
            markdown_content.append(f"## {para_text}\n")
            current_heading_level = 2
        elif "标题3" in para_style:
            markdown_content.append(f"### {para_text}\n")
            current_heading_level = 3
        else:
            # 2. 处理普通段落中的数学公式（转换为LaTeX格式）
            processed_text = process_math_formulas(para_text)
            
            # 3. 识别特殊区块（如选择题、答案、解析）
            if para_text.startswith(("【答案】", "【解析】", "【分析】", "【详解】")):
                # 为特殊区块添加Markdown强调格式
                if "【答案】" in para_text:
                    processed_text = f"**{processed_text}**"  # 答案加粗
                elif "【解析】" in para_text:
                    processed_text = f"> {processed_text}"  # 解析用引用格式
                markdown_content.append(f"{processed_text}\n")
            elif re.match(r"^\d+\.", para_text):  # 识别题号（如"1."、"2."）
                markdown_content.append(f"#### {processed_text}\n")  # 题号用四级标题
            else:
                markdown_content.append(f"{processed_text}\n")

    # 4. 处理表格（Word中的表格转换为Markdown表格）
    for table in doc.tables:
        table_markdown = []
        # 提取表头
        header_cells = table.rows[0].cells
        header = "| " + " | ".join([process_math_formulas(cell.text.strip()) for cell in header_cells]) + " |"
        separator = "| " + " | ".join(["---" for _ in header_cells]) + " |"
        table_markdown.append(header)
        table_markdown.append(separator)
        # 提取表格内容
        for row in table.rows[1:]:
            row_cells = row.cells
            row_content = "| " + " | ".join([process_math_formulas(cell.text.strip()) for cell in row_cells]) + " |"
            table_markdown.append(row_content)
        # 将表格添加到Markdown内容中
        markdown_content.append("\n".join(table_markdown) + "\n")

    return "".join(markdown_content)

def process_math_formulas(text):
    """
    将普通文本中的数学表达式转换为LaTeX格式
    处理规则：
    - 根号：√x → \sqrt{x}
    - 平方/次方：x² → x^2，(a+b)³ → (a+b)^3
    - 分数：a/b → \frac{a}{b}
    - 三角函数：sin30° → \sin30^\circ
    - 科学计数法：3.74×10^8 → 3.74\times10^8
    - 不等式：≤ → \leq，≥ → \geq
    """
    # 1. 处理根号（如√2、√(a+b)）
    text = re.sub(r"√(\w+)", r"\\sqrt{\1}", text)  # 简单根号：√x → \sqrt{x}
    text = re.sub(r"√\((.+?)\)", r"\\sqrt{\1}", text)  # 带括号根号：√(a+b) → \sqrt{a+b}

    # 2. 处理平方/次方（如x²、(a+b)³）
    text = re.sub(r"(\w+)²", r"\1^2", text)  # x² → x^2
    text = re.sub(r"(\w+)³", r"\1^3", text)  # x³ → x^3
    text = re.sub(r"\((.+?)\)²", r"(\1)^2", text)  # (a+b)² → (a+b)^2
    text = re.sub(r"\((.+?)\)³", r"(\1)^3", text)  # (a+b)³ → (a+b)^3

    # 3. 处理分数（如a/b、(a+b)/(c+d)）
    text = re.sub(r"(\w+)/(\w+)", r"\\frac{\1}{\2}", text)  # a/b → \frac{a}{b}
    text = re.sub(r"\((.+?)\)/\((.+?)\)", r"\\frac{\1}{\2}", text)  # (a+b)/(c+d) → \frac{a+b}{c+d}

    # 4. 处理三角函数（sin、cos、tan）
    text = re.sub(r"sin", r"\\sin", text)  # sin → \sin
    text = re.sub(r"cos", r"\\cos", text)  # cos → \cos
    text = re.sub(r"tan", r"\\tan", text)  # tan → \tan
    text = re.sub(r"(\d+)°", r"\1^\\circ", text)  # 30° → 30^\circ

    # 5. 处理科学计数法（3.74×10^8 → 3.74\times10^8）
    text = re.sub(r"×10\^", r"\\times10^", text)

    # 6. 处理不等式符号（≤、≥）
    text = re.sub(r"≤", r"\\leq", text)
    text = re.sub(r"≥", r"\\geq", text)

    # 7. 为所有数学公式添加$包裹（LaTeX行内公式）
    # 规则：包含LaTeX关键字（\sqrt、\frac、\sin等）的内容包裹为$...$
    text = re.sub(r"(\\sqrt{.+?}|\\frac{.+?}|\\sin.+?|\\cos.+?|\\tan.+?|\w+\^\\circ|\d+\\times10\^\d+)", r"$\1$", text)

    return text

def save_markdown(content, output_path):
    """
    将处理后的内容保存为Markdown文件
    :param content: 处理后的Markdown文本
    :param output_path: 输出文件路径（如"青岛中考数学真题.md"）
    """
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Markdown文件已保存至：{output_path}")

if __name__ == "__main__":
    # 1. 配置文件路径
    word_path = "Math/精品解析：2025年山东省青岛市中考数学真题（解析版）.docx"  # 输入Word文件路径
    output_path = "青岛中考数学真题_带LaTeX.md"  # 输出Markdown文件路径

    # 2. 验证Word文件是否存在
    if not Path(word_path).exists():
        print(f"错误：Word文件不存在，请检查路径：{word_path}")
        exit(1)

    # 3. 读取Word文档
    doc = Document(word_path)
    # 设置中文字体支持（避免乱码）
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 4. 提取并处理内容（转换为带LaTeX的Markdown）
    markdown_content = extract_text_with_format(doc)

    # 5. 保存结果
    save_markdown(markdown_content, output_path)